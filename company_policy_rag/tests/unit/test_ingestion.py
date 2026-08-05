from __future__ import annotations

import json
from pathlib import Path
import pytest

from backend.models.document import DocumentCategory, DocumentMetadata, DocumentType, RawDocument
from backend.models.chunk import Chunk, ChunkMetadata, ChunkRole, ContentType
from backend.models.ingestion import IngestionRequest, IngestionResult

from backend.ingestion.loaders.base import BaseLoader
from backend.ingestion.loaders.pdf import PDFLoader
from backend.ingestion.loaders.docx import DocxLoader
from backend.ingestion.loaders.txt import TxtLoader
from backend.ingestion.loaders.markdown import MarkdownLoader
from backend.ingestion.loaders.html import HTMLLoader
from backend.ingestion.loaders.csv import CSVLoader
from backend.ingestion.loaders.json import JSONLoader
from backend.ingestion.loaders.loader_factory import LoaderFactory, get_loader_for_file, load_document

from backend.ingestion.chunkers.base import BaseChunker
from backend.ingestion.chunkers.recursive import RecursiveChunker
from backend.ingestion.chunkers.semantic import SemanticChunker
from backend.ingestion.chunkers.markdown_aware import MarkdownAwareChunker
from backend.ingestion.chunkers.heading_aware import HeadingAwareChunker
from backend.ingestion.chunkers.table_aware import TableAwareChunker
from backend.ingestion.chunkers.adaptive_chunker import AdaptiveChunker

from backend.utils.hashing import compute_file_hash, compute_string_hash
from backend.utils.section_tracker import SectionTracker, parse_section_heading


# ── Loader Tests ─────────────────────────────────────────────────────────────

def test_load_txt_document(tmp_path: Path):
    txt_file = tmp_path / "handbook.txt"
    txt_content = (
        "I. GENERAL POLICIES\n"
        "A. At-Will Employment\n"
        "Employment at Company is on an at-will basis.\n\n"
        "B. Equal Opportunity\n"
        "We are an equal opportunity employer.\n"
    )
    txt_file.write_text(txt_content, encoding="utf-8")

    loader = TxtLoader()
    assert loader.supports(txt_file)

    docs = loader.load(txt_file, base_metadata={"category": "policy"})
    assert len(docs) == 1
    doc = docs[0]
    assert doc.metadata.document_type == DocumentType.TXT
    assert doc.metadata.category == "policy"
    assert doc.metadata.source_file == "handbook.txt"
    assert doc.metadata.section_path is not None
    assert "At-Will Employment" in doc.metadata.section_path or "GENERAL POLICIES" in doc.metadata.section_path
    assert "at-will basis" in doc.content


def test_load_markdown_document(tmp_path: Path):
    md_file = tmp_path / "guide.md"
    md_content = (
        "# Engineering Guidebook\n\n"
        "## Setup Instructions\n"
        "Run the following command:\n\n"
        "```bash\n"
        "pip install -r requirements.txt\n"
        "```\n\n"
        "## Support Matrix\n"
        "| Component | Status |\n"
        "|---|---|\n"
        "| Backend | Ready |\n"
        "| Frontend | Ready |\n"
    )
    md_file.write_text(md_content, encoding="utf-8")

    loader = MarkdownLoader()
    assert loader.supports(md_file)

    docs = loader.load(md_file)
    assert len(docs) == 1
    doc = docs[0]
    assert doc.metadata.document_type == DocumentType.MARKDOWN
    assert doc.metadata.has_code is True
    assert doc.metadata.has_tables is True
    assert doc.metadata.section_title is not None


def test_load_html_document(tmp_path: Path):
    html_file = tmp_path / "page.html"
    html_content = (
        "<html><head><title>Company Portal</title></head><body>"
        "<h1>Employee Code of Conduct</h1>"
        "<h2>Article 1: Respect</h2>"
        "<p>Treat everyone with dignity.</p>"
        "<table><tr><th>Role</th><th>Permission</th></tr>"
        "<tr><td>Admin</td><td>Full</td></tr></table>"
        "<code>console.log('hello');</code>"
        "</body></html>"
    )
    html_file.write_text(html_content, encoding="utf-8")

    loader = HTMLLoader()
    assert loader.supports(html_file)

    docs = loader.load(html_file)
    assert len(docs) == 1
    doc = docs[0]
    assert doc.metadata.document_type == DocumentType.HTML
    assert doc.metadata.section_title == "Company Portal"
    assert doc.metadata.has_tables is True
    assert doc.metadata.has_code is True
    assert "| Role | Permission |" in doc.content


def test_load_csv_document(tmp_path: Path):
    csv_file = tmp_path / "employees.csv"
    csv_content = (
        "ID,Name,Department,Role\n"
        "1,Alice,Engineering,Lead\n"
        "2,Bob,HR,Manager\n"
        "3,Charlie,Finance,Analyst\n"
    )
    csv_file.write_text(csv_content, encoding="utf-8")

    loader = CSVLoader()
    assert loader.supports(csv_file)

    docs = loader.load(csv_file)
    assert len(docs) == 1
    doc = docs[0]
    assert doc.metadata.document_type == DocumentType.CSV
    assert doc.metadata.has_tables is True
    assert doc.metadata.extra["columns"] == ["ID", "Name", "Department", "Role"]
    assert doc.metadata.extra["row_count"] == 3
    assert "| ID | Name | Department | Role |" in doc.content


def test_load_json_document(tmp_path: Path):
    json_file = tmp_path / "config.json"
    json_data = {"app_name": "RAG Chatbot", "version": "1.0.0", "features": ["auth", "search"]}
    json_file.write_text(json.dumps(json_data), encoding="utf-8")

    loader = JSONLoader()
    assert loader.supports(json_file)

    docs = loader.load(json_file)
    assert len(docs) == 1
    doc = docs[0]
    assert doc.metadata.document_type == DocumentType.JSON
    assert "RAG Chatbot" in doc.content

    # Test JSONL format
    jsonl_file = tmp_path / "data.jsonl"
    jsonl_lines = [
        json.dumps({"id": 1, "topic": "Security"}),
        json.dumps({"id": 2, "topic": "Compliance"}),
    ]
    jsonl_file.write_text("\n".join(jsonl_lines), encoding="utf-8")

    jsonl_docs = loader.load(jsonl_file)
    assert len(jsonl_docs) == 2
    assert jsonl_docs[0].metadata.extra["record_index"] == 1
    assert "Security" in jsonl_docs[0].content


def test_load_docx_document(tmp_path: Path):
    import docx
    docx_file = tmp_path / "sample.docx"
    doc_obj = docx.Document()
    doc_obj.add_heading("I. Security Policy", level=1)
    doc_obj.add_paragraph("All employees must use MFA.")
    
    # Add a table
    table = doc_obj.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Level"
    table.cell(0, 1).text = "Requirement"
    table.cell(1, 0).text = "High"
    table.cell(1, 1).text = "Hardware Key"
    
    doc_obj.save(str(docx_file))

    loader = DocxLoader()
    assert loader.supports(docx_file)

    docs = loader.load(docx_file)
    assert len(docs) == 1
    doc = docs[0]
    assert doc.metadata.document_type == DocumentType.DOCX
    assert doc.metadata.has_tables is True
    assert "Security Policy" in doc.content
    assert "MFA" in doc.content


def test_load_pdf_document(tmp_path: Path):
    import fitz
    pdf_file = tmp_path / "policy.pdf"
    pdf_doc = fitz.open()
    page1 = pdf_doc.new_page()
    page1.insert_text((50, 50), "I. OVERVIEW\nThis is page 1 of company policy.")
    page2 = pdf_doc.new_page()
    page2.insert_text((50, 50), "A. Scope\nThis is page 2 covering scope.")
    pdf_doc.save(str(pdf_file))
    pdf_doc.close()

    loader = PDFLoader()
    assert loader.supports(pdf_file)

    docs = loader.load(pdf_file)
    assert len(docs) == 2
    assert docs[0].metadata.document_type == DocumentType.PDF
    assert docs[0].metadata.page_number == 1
    assert docs[0].metadata.total_pages == 2
    assert "page 1" in docs[0].content
    assert docs[1].metadata.page_number == 2
    assert "page 2" in docs[1].content


def test_loader_factory(tmp_path: Path):
    f_txt = tmp_path / "doc.txt"
    f_txt.write_text("Hello", encoding="utf-8")
    
    factory = LoaderFactory()
    loader = factory.get_loader_for_file(f_txt)
    assert isinstance(loader, TxtLoader)

    docs = load_document(f_txt)
    assert len(docs) == 1
    assert docs[0].content == "Hello"


# ── Chunker Tests ────────────────────────────────────────────────────────────

def test_recursive_chunker():
    doc = RawDocument(
        content="Paragraph 1 text.\n\nParagraph 2 text.\n\nParagraph 3 text is longer to demonstrate recursive splitting capability.",
        metadata=DocumentMetadata(
            source_file="test.txt",
            file_path="test.txt",
            file_hash="1234567890abcdef",
            document_type=DocumentType.TXT,
        ),
    )

    chunker = RecursiveChunker(chunk_size=50, chunk_overlap=10)
    chunks = chunker.chunk([doc])

    assert len(chunks) > 0
    for c in chunks:
        assert isinstance(c, Chunk)
        assert c.metadata.chunk_strategy == "recursive"
        assert c.token_count > 0


def test_semantic_chunker():
    doc = RawDocument(
        content=(
            "First sentence about security. Second sentence detailing firewall rules. "
            "Third sentence regarding access controls.\n\n"
            "New paragraph covering data privacy. Compliance guidelines follow closely."
        ),
        metadata=DocumentMetadata(
            source_file="test.txt",
            file_path="test.txt",
            file_hash="1234567890abcdef",
            document_type=DocumentType.TXT,
        ),
    )

    chunker = SemanticChunker(chunk_size=30, chunk_overlap=10)
    chunks = chunker.chunk([doc])

    assert len(chunks) > 0
    assert chunks[0].metadata.chunk_strategy == "semantic"
    assert chunks[0].metadata.content_type == ContentType.PROSE


def test_markdown_aware_chunker():
    doc = RawDocument(
        content=(
            "# Header 1\n"
            "Some introduction prose.\n\n"
            "```python\n"
            "def calculate_total(price, tax):\n"
            "    return price * (1 + tax)\n"
            "```\n\n"
            "## Header 2\n"
            "More explanation after code block."
        ),
        metadata=DocumentMetadata(
            source_file="test.md",
            file_path="test.md",
            file_hash="1234567890abcdef",
            document_type=DocumentType.MARKDOWN,
            has_code=True,
        ),
    )

    chunker = MarkdownAwareChunker()
    chunks = chunker.chunk([doc])

    code_chunks = [c for c in chunks if c.metadata.content_type == ContentType.CODE]
    assert len(code_chunks) == 1
    assert code_chunks[0].metadata.is_atomic is True
    assert "def calculate_total" in code_chunks[0].text


def test_heading_aware_chunker():
    doc = RawDocument(
        content=(
            "I. OVERVIEW\n"
            "General company overview.\n"
            "A. Scope of Policy\n"
            "This policy applies to all employees.\n"
            "1.1 Remote Work\n"
            "Remote work requires manager approval."
        ),
        metadata=DocumentMetadata(
            source_file="policy.txt",
            file_path="policy.txt",
            file_hash="1234567890abcdef",
            document_type=DocumentType.TXT,
        ),
    )

    chunker = HeadingAwareChunker()
    chunks = chunker.chunk([doc])

    assert len(chunks) > 0
    assert chunks[0].metadata.chunk_strategy == "heading_aware"
    # Verify section_path metadata propagation
    section_paths = [c.metadata.section_path for c in chunks if c.metadata.section_path]
    assert len(section_paths) > 0


def test_table_aware_chunker():
    doc = RawDocument(
        content=(
            "Intro text before table.\n\n"
            "| Tier | Discount | Criteria |\n"
            "|---|---|---|\n"
            "| Bronze | 5% | Standard |\n"
            "| Silver | 10% | Premier |\n"
            "| Gold | 15% | Enterprise |\n\n"
            "Outro text after table."
        ),
        metadata=DocumentMetadata(
            source_file="pricing.md",
            file_path="pricing.md",
            file_hash="1234567890abcdef",
            document_type=DocumentType.MARKDOWN,
            has_tables=True,
        ),
    )

    chunker = TableAwareChunker(chunk_size=100)
    chunks = chunker.chunk([doc])

    table_chunks = [c for c in chunks if c.metadata.content_type == ContentType.TABLE]
    assert len(table_chunks) >= 1
    assert table_chunks[0].metadata.is_atomic is True
    assert "| Tier | Discount | Criteria |" in table_chunks[0].text


def test_adaptive_chunker():
    doc_table = RawDocument(
        content="| Col1 | Col2 |\n|---|---|\n| Val1 | Val2 |",
        metadata=DocumentMetadata(
            source_file="table.md",
            file_path="table.md",
            file_hash="1111111111111111",
            document_type=DocumentType.MARKDOWN,
            has_tables=True,
        ),
    )
    doc_code = RawDocument(
        content="```python\nprint('hello')\n```",
        metadata=DocumentMetadata(
            source_file="code.py",
            file_path="code.py",
            file_hash="2222222222222222",
            document_type=DocumentType.MARKDOWN,
            has_code=True,
        ),
    )

    adaptive = AdaptiveChunker()

    selected_table = adaptive.select_chunker(doc_table)
    assert isinstance(selected_table, TableAwareChunker)

    selected_code = adaptive.select_chunker(doc_code)
    assert isinstance(selected_code, MarkdownAwareChunker)

    chunks = adaptive.chunk([doc_table, doc_code])
    assert len(chunks) >= 2


# ── Hashing Tests ────────────────────────────────────────────────────────────

def test_hashing_utils(tmp_path: Path):
    f = tmp_path / "sample.txt"
    f.write_text("Hello World Hash Test", encoding="utf-8")

    file_hash = compute_file_hash(f)
    assert len(file_hash) == 16

    str_hash = compute_string_hash("Hello World Hash Test")
    assert file_hash == str_hash
