from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

import zipfile

from backend.ingestion.loaders.base import BaseLoader
from backend.models.document import DocumentMetadata, DocumentType, RawDocument
from backend.utils.logging import logger
from backend.utils.section_tracker import SectionTracker, clean_title


class DocxLoader(BaseLoader):
    """Loader for Microsoft Word (.docx) documents."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".docx", ".doc"]

    def load(
        self,
        file_path: Path,
        base_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[RawDocument]:
        base_meta = self._build_base_metadata(file_path, DocumentType.DOCX, base_metadata)

        try:
            import docx
            import docx.opc.exceptions
        except ImportError as e:
            logger.error("python-docx is not installed: %s", e)
            raise RuntimeError("python-docx required for docx files") from e


        try:
            doc = docx.Document(str(file_path))
        except (Exception, zipfile.BadZipFile, docx.opc.exceptions.PackageNotFoundError, KeyError) as e:
            logger.warning("Failed to parse DOCX file %s: %s", file_path, e)
            return [RawDocument(content="", metadata=base_meta)]

        section_tracker = SectionTracker()
        content_parts: List[str] = []
        has_tables = False
        has_code = False

        # Process document elements in order: paragraphs and tables
        for element in doc.element.body:
            tag_name = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag_name == "p":
                # Paragraph element
                p_text = "".join(node.text for node in element.iter() if node.text).strip()
                if not p_text:
                    continue

                # Check if paragraph has a Heading style
                style_name = ""
                for child in element:
                    if child.tag.endswith("pPr"):
                        for subchild in child:
                            if subchild.tag.endswith("pStyle"):
                                style_name = subchild.attrib.get(
                                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
                                )

                if "Heading" in style_name or "heading" in style_name.lower():
                    # Parse level
                    level_str = "".join(filter(str.isdigit, style_name))
                    level = int(level_str) if level_str else 1
                    formatted_heading = f"{'#' * min(6, level)} {p_text}"
                    content_parts.append(formatted_heading)
                    section_tracker.process_line(formatted_heading)
                else:
                    content_parts.append(p_text)
                    section_tracker.process_line(p_text)

            elif tag_name == "tbl":
                # Table element
                has_tables = True
                table_lines: List[str] = []
                for row in element.iter():
                    if row.tag.endswith("tr"):
                        cells = []
                        for cell in row.iter():
                            if cell.tag.endswith("tc"):
                                cell_text = "".join(node.text for node in cell.iter() if node.text).strip()
                                cells.append(cell_text.replace("\n", " "))
                        if cells:
                            table_lines.append("| " + " | ".join(cells) + " |")

                if table_lines:
                    # Insert header divider line if needed
                    header_cells = len(table_lines[0].split("|")) - 2
                    divider = "| " + " | ".join(["---"] * max(1, header_cells)) + " |"
                    table_lines.insert(1, divider)
                    table_md = "\n".join(table_lines)
                    content_parts.append(table_md)

        full_content = "\n\n".join(content_parts)
        ctx = section_tracker.current_context()

        final_meta = base_meta.model_copy(
            update={
                "section_title": ctx.section_title,
                "section_number": ctx.section_number,
                "section_path": ctx.section_path,
                "section_level": ctx.section_level,
                "has_tables": has_tables,
                "has_code": has_code,
            }
        )

        return [RawDocument(content=full_content, metadata=final_meta)]
